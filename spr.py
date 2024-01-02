import speech_recognition as sr
from docx import Document

r = sr.Recognizer()


harvard = sr.AudioFile('harvard.wav')
with harvard as source:
    audio = r.record(source)
a=r.recognize_google(audio)

document = Document()

document.add_heading('Speech Recognition', 0)

p = document.add_paragraph(a)
document.save('demo.docx')
