import speech_recognition as sr
from docx import Document

r = sr.Recognizer()

mic = sr.Microphone()

with mic as source:
    print("Say Something")
    audio = r.listen(source)
    print("Time Over")

b=r.recognize_google(audio)
document = Document()

document.add_heading('Speech Recognition', 0)

p = document.add_paragraph(b)
document.save('demo_mi_new.docx')
